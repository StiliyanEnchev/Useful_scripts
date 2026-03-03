import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from docx import Document

# === FILE PATHS ===
xml_file = "WordPress.2026-03-03.xml"
output_file = "exported_posts.docx"

# === LOAD XML ===
tree = ET.parse(xml_file)
root = tree.getroot()

# WordPress namespaces
namespaces = {
    'content': 'http://purl.org/rss/1.0/modules/content/',
    'wp': 'http://wordpress.org/export/1.2/'
}

document = Document()

# === LOOP THROUGH POSTS IN ORIGINAL ORDER ===
for item in root.findall(".//item"):
    post_type = item.find("wp:post_type", namespaces)
    status = item.find("wp:status", namespaces)

    # Only published posts
    if post_type is not None and post_type.text == "post" and status.text == "publish":
        title = item.find("title").text or ""
        content = item.find("content:encoded", namespaces).text or ""

        # Remove HTML tags
        clean_text = BeautifulSoup(content, "html.parser").get_text()

        # Add to document
        document.add_heading(title, level=1)
        document.add_paragraph(clean_text)
        document.add_page_break()

# Save
document.save(output_file)

print("Done. File saved as:", output_file)